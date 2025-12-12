"""
–°–∫—Ä–∏–ø—Ç –¥–ª—è –∏–Ω–¥–µ–∫—Å–∞—Ü–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –≤ –≤–µ–∫—Ç–æ—Ä–Ω—É—é –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö Chroma.
–ß–∏—Ç–∞–µ—Ç –¥–æ–∫—É–º–µ–Ω—Ç—ã –∏–∑ data/raw (PDF, TXT, DOCX, MD), —Ä–∞–∑–±–∏–≤–∞–µ—Ç –Ω–∞ —á–∞–Ω–∫–∏, 
–≥–µ–Ω–µ—Ä–∏—Ä—É–µ—Ç embeddings –∏ —Å–æ—Ö—Ä–∞–Ω—è–µ—Ç –≤ Chroma.
"""

import os
import yaml
import tiktoken
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter
try:
    from langchain_chroma import Chroma
except ImportError:
    # Fallback –¥–ª—è —Å—Ç–∞—Ä—ã—Ö –≤–µ—Ä—Å–∏–π
    from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document

# –ù–∞—Å—Ç—Ä–æ–π–∫–∞ –ª–æ–≥–∏—Ä–æ–≤–∞–Ω–∏—è
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# –ö–æ–Ω—Ñ–∏–≥—É—Ä–∞—Ü–∏—è
DATA_DIR = Path(__file__).parent.parent / "data" / "raw"  
PROCESSED_DIR = Path(__file__).parent.parent / "data" / "processed"
CHROMA_PERSIST_DIR = Path(__file__).parent.parent / "chroma_db"
COLLECTION_NAME = "engineer_bot"

# –ü–∞—Ä–∞–º–µ—Ç—Ä—ã —á–∞–Ω–∫–∏–Ω–≥–∞
CHUNK_SIZE = 500  # —Ç–æ–∫–µ–Ω–æ–≤
CHUNK_OVERLAP = 75  # 15% –æ—Ç 500

# –ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º—ã–µ —Ñ–æ—Ä–º–∞—Ç—ã
SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx', '.md'}


def load_metadata(file_path: Path) -> Dict[str, Any]:
    """–ó–∞–≥—Ä—É–∂–∞–µ—Ç –º–µ—Ç–∞–¥–∞–Ω–Ω—ã–µ –∏–∑ YAML —Ñ–∞–π–ª–∞."""
    # –ü—Ä–æ–±—É–µ–º –Ω–∞–π—Ç–∏ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã–µ –≤ processed –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏–∏
    metadata_file = PROCESSED_DIR / f"{file_path.stem}.metadata.yaml"
    if not metadata_file.exists():
        # –ü—Ä–æ–±—É–µ–º –≤ —Ç–æ–π –∂–µ –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏–∏
        metadata_file = file_path.parent / f"{file_path.stem}.metadata.yaml"
    
    if metadata_file.exists():
        try:
            with open(metadata_file, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f) or {}
        except Exception as e:
            logger.warning(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –∑–∞–≥—Ä—É–∑–∫–µ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö {metadata_file}: {e}")
    
    return {}


def count_tokens(text: str) -> int:
    """–ü–æ–¥—Å—á–∏—Ç—ã–≤–∞–µ—Ç –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Ç–æ–∫–µ–Ω–æ–≤ –≤ —Ç–µ–∫—Å—Ç–µ (–ø—Ä–∏–±–ª–∏–∑–∏—Ç–µ–ª—å–Ω–æ)."""
    encoding = tiktoken.get_encoding("cl100k_base")
    return len(encoding.encode(text))


def parse_pdf(file_path: Path) -> str:
    """–ü–∞—Ä—Å–∏—Ç PDF —Ñ–∞–π–ª –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Ç–µ–∫—Å—Ç."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –ø–∞—Ä—Å–∏–Ω–≥–µ PDF {file_path}: {e}")
        raise


def parse_docx(file_path: Path) -> str:
    """–ü–∞—Ä—Å–∏—Ç DOCX —Ñ–∞–π–ª –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Ç–µ–∫—Å—Ç."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –ø–∞—Ä—Å–∏–Ω–≥–µ DOCX {file_path}: {e}")
        raise


def parse_txt(file_path: Path) -> str:
    """–ü–∞—Ä—Å–∏—Ç TXT —Ñ–∞–π–ª –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Ç–µ–∫—Å—Ç."""
    try:
        # –ü—Ä–æ–±—É–µ–º —Ä–∞–∑–Ω—ã–µ –∫–æ–¥–∏—Ä–æ–≤–∫–∏
        encodings = ['utf-8', 'cp1251', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"–ù–µ —É–¥–∞–ª–æ—Å—å –¥–µ–∫–æ–¥–∏—Ä–æ–≤–∞—Ç—å —Ñ–∞–π–ª {file_path}")
    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –ø–∞—Ä—Å–∏–Ω–≥–µ TXT {file_path}: {e}")
        raise


def parse_markdown(file_path: Path) -> str:
    """–ü–∞—Ä—Å–∏—Ç Markdown —Ñ–∞–π–ª –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Ç–µ–∫—Å—Ç."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –ø–∞—Ä—Å–∏–Ω–≥–µ Markdown {file_path}: {e}")
        raise


def load_document_content(file_path: Path) -> str:
    """–ó–∞–≥—Ä—É–∂–∞–µ—Ç —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç –µ–≥–æ —Ç–∏–ø–∞."""
    suffix = file_path.suffix.lower()
    
    if suffix == '.pdf':
        return parse_pdf(file_path)
    elif suffix == '.docx':
        return parse_docx(file_path)
    elif suffix == '.txt':
        return parse_txt(file_path)
    elif suffix == '.md':
        return parse_markdown(file_path)
    else:
        raise ValueError(f"–ù–µ–ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º—ã–π —Ñ–æ—Ä–º–∞—Ç —Ñ–∞–π–ª–∞: {suffix}")


def load_documents() -> List[Document]:
    """–ó–∞–≥—Ä—É–∂–∞–µ—Ç –≤—Å–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã –∏–∑ data/raw —Å –º–µ—Ç–∞–¥–∞–Ω–Ω—ã–º–∏."""
    documents = []
    
    if not DATA_DIR.exists():
        logger.warning(f"–î–∏—Ä–µ–∫—Ç–æ—Ä–∏—è {DATA_DIR} –Ω–µ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç. –°–æ–∑–¥–∞—é...")
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        return documents
    
    # –ù–∞—Ö–æ–¥–∏–º –≤—Å–µ –ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º—ã–µ —Ñ–∞–π–ª—ã (—Ä–µ–∫—É—Ä—Å–∏–≤–Ω–æ, –≤–∫–ª—é—á–∞—è –ø–æ–¥–¥–∏—Ä–µ–∫—Ç–æ—Ä–∏–∏)
    all_files = []
    for ext in SUPPORTED_EXTENSIONS:
        # –ò—â–µ–º –≤ –∫–æ—Ä–Ω–µ –∏ –≤–æ –≤—Å–µ—Ö –ø–æ–¥–¥–∏—Ä–µ–∫—Ç–æ—Ä–∏—è—Ö
        all_files.extend(list(DATA_DIR.rglob(f"*{ext}")))
    
    if not all_files:
        logger.warning(f"–ù–µ –Ω–∞–π–¥–µ–Ω–æ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –≤ {DATA_DIR}")
        return documents
    
    logger.info(f"–ù–∞–π–¥–µ–Ω–æ {len(all_files)} –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏")
    
    for file_path in all_files:
        try:
            # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º —Ñ–∞–π–ª—ã –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö
            if file_path.name.endswith('.metadata.yaml'):
                continue
            
            logger.info(f"–û–±—Ä–∞–±–æ—Ç–∫–∞ —Ñ–∞–π–ª–∞: {file_path.name}")
            
            # –ó–∞–≥—Ä—É–∂–∞–µ–º —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ
            content = load_document_content(file_path)
            
            if not content.strip():
                logger.warning(f"–§–∞–π–ª {file_path.name} –ø—É—Å—Ç, –ø—Ä–æ–ø—É—Å–∫–∞—é")
                continue
            
            # –ó–∞–≥—Ä—É–∂–∞–µ–º –º–µ—Ç–∞–¥–∞–Ω–Ω—ã–µ
            metadata = load_metadata(file_path)
            
            # –°–æ–∑–¥–∞–µ–º –±–∞–∑–æ–≤—ã–π –¥–æ–∫—É–º–µ–Ω—Ç
            doc_metadata = {
                "source": file_path.name,
                "file_type": file_path.suffix.lower(),
                "type": metadata.get("type", "unknown"),
                "revision": metadata.get("revision", "unknown"),
                "date": metadata.get("date", "unknown"),
                "category": metadata.get("category", "unknown"),
                "name": metadata.get("name", file_path.stem),
            }
            
            # –î–æ–±–∞–≤–ª—è–µ–º —Ç–µ–≥–∏, –µ—Å–ª–∏ –µ—Å—Ç—å
            if "tags" in metadata:
                doc_metadata["tags"] = ", ".join(metadata["tags"])
            
            documents.append(Document(page_content=content, metadata=doc_metadata))
            logger.info(f"–ó–∞–≥—Ä—É–∂–µ–Ω –¥–æ–∫—É–º–µ–Ω—Ç: {file_path.name} ({len(content)} —Å–∏–º–≤–æ–ª–æ–≤)")
        
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –æ–±—Ä–∞–±–æ—Ç–∫–µ {file_path.name}: {e}")
            continue
    
    return documents


def chunk_documents(documents: List[Document], chunk_size: int, chunk_overlap: int) -> List[Document]:
    """
    –†–∞–∑–±–∏–≤–∞–µ—Ç –¥–æ–∫—É–º–µ–Ω—Ç—ã –Ω–∞ —á–∞–Ω–∫–∏ —Å —É—á–µ—Ç–æ–º —Ä–∞–∑–º–µ—Ä–∞ –≤ —Ç–æ–∫–µ–Ω–∞—Ö.
    
    Args:
        documents: –°–ø–∏—Å–æ–∫ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –¥–ª—è —Ä–∞–∑–±–∏–µ–Ω–∏—è
        chunk_size: –†–∞–∑–º–µ—Ä —á–∞–Ω–∫–∞ –≤ —Ç–æ–∫–µ–Ω–∞—Ö
        chunk_overlap: –ü–µ—Ä–µ–∫—Ä—ã—Ç–∏–µ —á–∞–Ω–∫–æ–≤ –≤ —Ç–æ–∫–µ–Ω–∞—Ö
    """
    
    # –ò—Å–ø–æ–ª—å–∑—É–µ–º RecursiveCharacterTextSplitter
    # –û–Ω –ø—ã—Ç–∞–µ—Ç—Å—è —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å —Å—Ç—Ä—É–∫—Ç—É—Ä—É –¥–æ–∫—É–º–µ–Ω—Ç–∞ (–ø–∞—Ä–∞–≥—Ä–∞—Ñ—ã, –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=count_tokens,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    chunked_docs = []
    
    for doc in documents:
        try:
            # –†–∞–∑–±–∏–≤–∞–µ–º –¥–æ–∫—É–º–µ–Ω—Ç –Ω–∞ —á–∞–Ω–∫–∏
            chunks = text_splitter.split_documents([doc])
            
            # –î–æ–±–∞–≤–ª—è–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ –Ω–æ–º–µ—Ä–µ —á–∞–Ω–∫–∞
            for i, chunk in enumerate(chunks):
                chunk.metadata["chunk_index"] = i
                chunk.metadata["total_chunks"] = len(chunks)
            
            chunked_docs.extend(chunks)
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Ä–∞–∑–±–∏–µ–Ω–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞ {doc.metadata.get('source', 'unknown')}: {e}")
            continue
    
    return chunked_docs


def create_embeddings(use_api: bool = False):
    """
    –°–æ–∑–¥–∞–µ—Ç –º–æ–¥–µ–ª—å –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ embeddings.
    
    Args:
        use_api: –ï—Å–ª–∏ True, –∏—Å–ø–æ–ª—å–∑—É–µ—Ç Hugging Face Inference API (—Ç—Ä–µ–±—É–µ—Ç HF_TOKEN)
                 –ï—Å–ª–∏ False, –∏—Å–ø–æ–ª—å–∑—É–µ—Ç –ª–æ–∫–∞–ª—å–Ω—É—é –º–æ–¥–µ–ª—å (–±–µ—Å–ø–ª–∞—Ç–Ω–æ, –±—ã—Å—Ç—Ä–µ–µ)
    """
    if use_api:
        # –ò—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ Hugging Face Inference API (—Ç—Ä–µ–±—É–µ—Ç HF_TOKEN)
        hf_token = os.getenv("HUGGINGFACE_API_TOKEN")
        if not hf_token:
            print("HUGGINGFACE_API_TOKEN –Ω–µ –Ω–∞–π–¥–µ–Ω, –∏—Å–ø–æ–ª—å–∑—É–µ–º –ª–æ–∫–∞–ª—å–Ω—É—é –º–æ–¥–µ–ª—å")
            use_api = False
    
    if use_api:
        # –ò—Å–ø–æ–ª—å–∑—É–µ–º Inference API –¥–ª—è embeddings
        model_name = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
        embeddings = HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
    else:
        # –õ–æ–∫–∞–ª—å–Ω–∞—è –º–æ–¥–µ–ª—å - –±–µ—Å–ø–ª–∞—Ç–Ω–æ –∏ –±—ã—Å—Ç—Ä–µ–µ –¥–ª—è –∏–Ω–¥–µ–∫—Å–∞—Ü–∏–∏
        # –ò—Å–ø–æ–ª—å–∑—É–µ–º ruBERT –¥–ª—è –ª—É—á—à–µ–π —Ä–∞–±–æ—Ç—ã —Å —Ä—É—Å—Å–∫–∏–º —è–∑—ã–∫–æ–º
        model_name = "cointegrated/rubert-tiny2"
        embeddings = HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={'device': 'cpu'},  # –ò—Å–ø–æ–ª—å–∑—É–µ–º CPU, –º–æ–∂–Ω–æ –∏–∑–º–µ–Ω–∏—Ç—å –Ω–∞ 'cuda' –µ—Å–ª–∏ –µ—Å—Ç—å GPU
            encode_kwargs={'normalize_embeddings': True}  # –ù–æ—Ä–º–∞–ª–∏–∑–∞—Ü–∏—è –¥–ª—è cosine similarity
        )
    
    return embeddings


def ingest_documents(chunk_size: Optional[int] = None, chunk_overlap: Optional[int] = None):
    """
    –û—Å–Ω–æ–≤–Ω–∞—è —Ñ—É–Ω–∫—Ü–∏—è –∏–Ω–¥–µ–∫—Å–∞—Ü–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤.
    
    Args:
        chunk_size: –†–∞–∑–º–µ—Ä —á–∞–Ω–∫–∞ –≤ —Ç–æ–∫–µ–Ω–∞—Ö (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é CHUNK_SIZE)
        chunk_overlap: –ü–µ—Ä–µ–∫—Ä—ã—Ç–∏–µ —á–∞–Ω–∫–æ–≤ –≤ —Ç–æ–∫–µ–Ω–∞—Ö (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é CHUNK_OVERLAP)
    """
    logger.info("üöÄ –ù–∞—á–∞–ª–æ –∏–Ω–¥–µ–∫—Å–∞—Ü–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤...")
    
    # –ò—Å–ø–æ–ª—å–∑—É–µ–º –ø–µ—Ä–µ–¥–∞–Ω–Ω—ã–µ –ø–∞—Ä–∞–º–µ—Ç—Ä—ã –∏–ª–∏ –∑–Ω–∞—á–µ–Ω–∏—è –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é
    chunk_size = chunk_size or CHUNK_SIZE
    chunk_overlap = chunk_overlap or CHUNK_OVERLAP
    
    # 1. –ó–∞–≥—Ä—É–∂–∞–µ–º –¥–æ–∫—É–º–µ–Ω—Ç—ã
    logger.info(f"–ó–∞–≥—Ä—É–∑–∫–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –∏–∑ {DATA_DIR}...")
    documents = load_documents()
    
    if not documents:
        logger.error("–ù–µ –Ω–∞–π–¥–µ–Ω–æ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –¥–ª—è –∏–Ω–¥–µ–∫—Å–∞—Ü–∏–∏")
        logger.info(f"–ü–æ–º–µ—Å—Ç–∏—Ç–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã (PDF, TXT, DOCX, MD) –≤ –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏—é {DATA_DIR}")
        return None
    
    logger.info(f"–ó–∞–≥—Ä—É–∂–µ–Ω–æ {len(documents)} –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤")
    
    # 2. –†–∞–∑–±–∏–≤–∞–µ–º –Ω–∞ —á–∞–Ω–∫–∏
    logger.info(f"–†–∞–∑–±–∏–µ–Ω–∏–µ –Ω–∞ —á–∞–Ω–∫–∏ (—Ä–∞–∑–º–µ—Ä: {chunk_size} —Ç–æ–∫–µ–Ω–æ–≤, overlap: {chunk_overlap})...")
    try:
        chunked_docs = chunk_documents(documents, chunk_size, chunk_overlap)
        logger.info(f"–°–æ–∑–¥–∞–Ω–æ {len(chunked_docs)} —á–∞–Ω–∫–æ–≤")
    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Ä–∞–∑–±–∏–µ–Ω–∏–∏ –Ω–∞ —á–∞–Ω–∫–∏: {e}")
        raise
    
    # 3. –°–æ–∑–¥–∞–µ–º –º–æ–¥–µ–ª—å embeddings
    logger.info("–ó–∞–≥—Ä—É–∑–∫–∞ –º–æ–¥–µ–ª–∏ –¥–ª—è embeddings...")
    try:
        embeddings = create_embeddings()
        logger.info("–ú–æ–¥–µ–ª—å –∑–∞–≥—Ä—É–∂–µ–Ω–∞")
    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –∑–∞–≥—Ä—É–∑–∫–µ –º–æ–¥–µ–ª–∏ embeddings: {e}")
        raise
    
    # 4. –°–æ–∑–¥–∞–µ–º –∏–ª–∏ –æ–±–Ω–æ–≤–ª—è–µ–º –≤–µ–∫—Ç–æ—Ä–Ω—É—é –±–∞–∑—É
    logger.info(f"–°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –≤ Chroma (–∫–æ–ª–ª–µ–∫—Ü–∏—è: {COLLECTION_NAME})...")
    
    try:
        # –£–¥–∞–ª—è–µ–º —Å—Ç–∞—Ä—É—é –±–∞–∑—É, –µ—Å–ª–∏ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç
        if CHROMA_PERSIST_DIR.exists():
            import shutil
            shutil.rmtree(CHROMA_PERSIST_DIR)
            logger.info("–£–¥–∞–ª–µ–Ω–∞ —Å—Ç–∞—Ä–∞—è –±–∞–∑–∞ –¥–∞–Ω–Ω—ã—Ö")
        
        # –°–æ–∑–¥–∞–µ–º –Ω–æ–≤—É—é –±–∞–∑—É
        vectorstore = Chroma.from_documents(
            documents=chunked_docs,
            embedding=embeddings,
            persist_directory=str(CHROMA_PERSIST_DIR),
            collection_name=COLLECTION_NAME
        )
        
        logger.info(f"–î–æ–∫—É–º–µ–Ω—Ç—ã –ø—Ä–æ–∏–Ω–¥–µ–∫—Å–∏—Ä–æ–≤–∞–Ω—ã –∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω—ã –≤ {CHROMA_PERSIST_DIR}")
        
        # 5. –ü—Ä–æ–≤–µ—Ä—è–µ–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –≤ –±–∞–∑–µ
        try:
            # –ü—ã—Ç–∞–µ–º—Å—è –ø–æ–ª—É—á–∏—Ç—å –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —á–µ—Ä–µ–∑ –ø—É–±–ª–∏—á–Ω—ã–π API
            if hasattr(vectorstore, '_collection'):
                count = vectorstore._collection.count()
            else:
                # –ê–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω—ã–π —Å–ø–æ—Å–æ–± - —á–µ—Ä–µ–∑ retriever
                count = len(chunked_docs)
            logger.info(f"–í—Å–µ–≥–æ –≤–µ–∫—Ç–æ—Ä–æ–≤ –≤ –±–∞–∑–µ: {count}")
        except Exception as e:
            logger.warning(f"–ù–µ —É–¥–∞–ª–æ—Å—å –ø–æ–ª—É—á–∏—Ç—å —Ç–æ—á–Ω–æ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –≤–µ–∫—Ç–æ—Ä–æ–≤: {e}")
            logger.info(f"–ò–Ω–¥–µ–∫—Å–∏—Ä–æ–≤–∞–Ω–æ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤: {len(chunked_docs)}")
        
        return vectorstore
    
    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ –≤ Chroma: {e}")
        raise


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="–ò–Ω–¥–µ–∫—Å–∞—Ü–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –≤ –≤–µ–∫—Ç–æ—Ä–Ω—É—é –±–∞–∑—É")
    parser.add_argument("--chunk-size", type=int, default=CHUNK_SIZE, help="–†–∞–∑–º–µ—Ä —á–∞–Ω–∫–∞ –≤ —Ç–æ–∫–µ–Ω–∞—Ö")
    parser.add_argument("--chunk-overlap", type=int, default=CHUNK_OVERLAP, help="–ü–µ—Ä–µ–∫—Ä—ã—Ç–∏–µ —á–∞–Ω–∫–æ–≤ –≤ —Ç–æ–∫–µ–Ω–∞—Ö")
    
    args = parser.parse_args()
    
    try:
        vectorstore = ingest_documents(
            chunk_size=args.chunk_size,
            chunk_overlap=args.chunk_overlap
        )
        if vectorstore:
            logger.info("\n–ò–Ω–¥–µ–∫—Å–∞—Ü–∏—è –∑–∞–≤–µ—Ä—à–µ–Ω–∞ —É—Å–ø–µ—à–Ω–æ!")
            logger.info(f"–ë–∞–∑–∞ –¥–∞–Ω–Ω—ã—Ö: {CHROMA_PERSIST_DIR}")
            logger.info(f"–ö–æ–ª–ª–µ–∫—Ü–∏—è: {COLLECTION_NAME}")
    except Exception as e:
        logger.error(f"\n–û—à–∏–±–∫–∞ –ø—Ä–∏ –∏–Ω–¥–µ–∫—Å–∞—Ü–∏–∏: {e}")
        import traceback
        traceback.print_exc()
        exit(1)

